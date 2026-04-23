def extract_text(filename: str, binary: bytes) -> tuple[str, list[str], list[dict]]:
    """
    Extract text content based on file type (Pipeline Phase 1).

    PDF files use dual-path fusion + layout reconstruction + line indexing.
    Other formats fall back to simple text extraction.

    Args:
        filename: File name
        binary: File binary content
    Returns:
        (indexed_text, lines, line_positions) tuple:
        - indexed_text: Text with line number indices
        - lines: List of original line texts
        - line_positions: List of per-line coordinate info (empty list for non-PDF formats)
    """
    fname_lower = filename.lower()

    try:
        if fname_lower.endswith(".pdf"):
            # Dual-path extraction
            meta_blocks = _extract_metadata_text(binary)
            ocr_blocks = []

            # Determine whether OCR supplementation is needed:
            # 1. Metadata text too short (< 100 chars)
            # 2. High garbled text ratio in metadata (caused by custom font mapping)
            meta_text_len = sum(len(b["text"]) for b in meta_blocks)
            need_ocr = False

            if meta_text_len < 100:
                logger.info("PDF metadata text too short, enabling OCR supplementation")
                need_ocr = True
            else:
                # Check metadata text quality: calculate valid line ratio
                # If many lines are judged as garbled by _is_valid_line, the PDF font mapping has issues
                valid_line_count = 0
                total_line_count = 0
                for b in meta_blocks:
                    text = b.get("text", "").strip()
                    if not text:
                        continue
                    total_line_count += 1
                    if _is_valid_line(text):
                        valid_line_count += 1
                if total_line_count > 0:
                    valid_ratio = valid_line_count / total_line_count
                    if valid_ratio < 0.6:
                        logger.info(
                            f"PDF metadata text quality low (valid line ratio {valid_ratio:.1%}), enabling OCR supplementation"
                        )
                        need_ocr = True

            if need_ocr:
                # Blackout strategy: black out metadata-extracted regions before OCR
                ocr_blocks = _extract_ocr_text(binary, meta_blocks=meta_blocks)

            # Text fusion
            fused_blocks = _fuse_text_blocks(meta_blocks, ocr_blocks)

            # Layout-aware sorting (prefer YOLOv10 layout detection, fall back to heuristic on failure)
            sorted_blocks = _layout_detect_reorder(fused_blocks, binary)

            # Build line-indexed text (with coordinate info)
            return _build_indexed_text(sorted_blocks)

        elif fname_lower.endswith(".docx"):
            from docx import Document
            doc = Document(BytesIO(binary))
            lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

            # Extract table content from DOCX
            # Reference: table handling in naive.py Docx class
            # Many resumes use table layouts for personal info; iterating only paragraphs would miss this content
            for table in doc.tables:
                for row in table.rows:
                    cells = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            cells.append(cell_text)
                    if not cells:
                        continue
                    row_text = " | ".join(cells)
                    # Deduplicate: skip if this row text already exists in lines
                    if row_text not in lines:
                        lines.append(row_text)

            indexed = "\n".join(f"[{i}]: {line}" for i, line in enumerate(lines))
            # DOCX has no coordinate info, return empty list
            return indexed, lines, []

        else:
            text = get_text(filename, binary)
            lines = [line.strip() for line in text.split("\n") if line.strip()]
            indexed = "\n".join(f"[{i}]: {line}" for i, line in enumerate(lines))
            return indexed, lines, []

    except Exception:
        logger.exception(f"Text extraction failed: {filename}")
        return "", [], []