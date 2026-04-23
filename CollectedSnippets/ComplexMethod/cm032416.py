def _decorate_docx(self, file_path: str) -> tuple[str, bytes]:
        from docx import Document
        from docx.enum.text import WD_TAB_ALIGNMENT
        from docx.shared import Pt

        document = Document(file_path)
        default_section = Document().sections[0]
        h1_size, h2_size, h3_size = self._get_heading_sizes()

        style_map = {
            "Normal": int(self._param.font_size),
            "Heading 1": h1_size,
            "Heading 2": h2_size,
            "Heading 3": h3_size,
        }
        for style_name, size in style_map.items():
            try:
                document.styles[style_name].font.size = Pt(size)
            except Exception:
                continue

        for section in document.sections:
            self._normalize_docx_section_geometry(section, default_section)
            available_width = self._get_docx_available_width(section)

            header = section.header
            header.is_linked_to_previous = False
            self._clear_docx_container(header)
            if self._param.header_text:
                paragraph = header.add_paragraph()
                paragraph.add_run(self._param.header_text)

            self._add_docx_watermark(section)

            footer = section.footer
            footer.is_linked_to_previous = False
            self._clear_docx_container(footer)
            if any(
                [
                    self._param.footer_text,
                    self._param.add_timestamp,
                    self._param.add_page_numbers,
                ]
            ):
                paragraph = footer.add_paragraph()
                paragraph.paragraph_format.tab_stops.add_tab_stop(
                    int(available_width // 2),
                    WD_TAB_ALIGNMENT.CENTER,
                )
                paragraph.paragraph_format.tab_stops.add_tab_stop(
                    int(available_width),
                    WD_TAB_ALIGNMENT.RIGHT,
                )

                if self._param.footer_text:
                    paragraph.add_run(self._param.footer_text)

                if self._param.add_timestamp or self._param.add_page_numbers:
                    paragraph.add_run("\t")

                if self._param.add_timestamp:
                    paragraph.add_run(self._get_timestamp_text())

                if self._param.add_page_numbers:
                    paragraph.add_run("\t")
                    self._append_docx_field(paragraph.add_run(), " PAGE ")

        document.save(file_path)
        with open(file_path, "rb") as f:
            file_bytes = f.read()
        return file_path, file_bytes